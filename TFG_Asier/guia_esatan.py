from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document()

# Page setup A4
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(20)
section.bottom_margin = Mm(18)
section.left_margin = Mm(22)
section.right_margin = Mm(18)
section.header_distance = Mm(8)
section.footer_distance = Mm(8)

# Base styles
styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal'].font.size = Pt(10.5)
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
styles['Title'].font.name = 'Calibri'
styles['Title'].font.size = Pt(22)
styles['Title'].font.bold = True
styles['Title']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

for sname, size, color in [('Heading 1', 15, '1F4E79'), ('Heading 2', 12.5, '1F4E79'), ('Heading 3', 11.2, '2F2F2F')]:
    st = styles[sname]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.bold = True
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    # color
    rpr = st.element.get_or_add_rPr()
    c = rpr.find(qn('w:color'))
    if c is None:
        c = OxmlElement('w:color')
        rpr.append(c)
    c.set(qn('w:val'), color)
    
# custom styles
if 'Small Table' not in styles:
    st = styles.add_style('Small Table', WD_STYLE_TYPE.PARAGRAPH)
    st.font.name = 'Calibri'
    st.font.size = Pt(8.5)
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

if 'Callout' not in styles:
    st = styles.add_style('Callout', WD_STYLE_TYPE.PARAGRAPH)
    st.font.name = 'Calibri'
    st.font.size = Pt(10)
    st.font.italic = False
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_text(cell, text, bold=False, style='Normal', align=None, size=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if style:
        p.style = styles[style]
    if align is not None:
        p.alignment = align
    r = p.add_run('' if text is None else str(text))
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    r.font.name = 'Calibri'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

def add_table(data, headers=None, style_small=False, col_widths=None, header_fill='D9EAF7'):
    rows = len(data) + (1 if headers else 0)
    cols = len(headers) if headers else len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    row_idx = 0
    if headers:
        for j,h in enumerate(headers):
            set_cell_text(table.rows[0].cells[j], h, bold=True, style='Small Table' if style_small else 'Normal', align=WD_ALIGN_PARAGRAPH.CENTER, size=8.5 if style_small else 10)
            shade_cell(table.rows[0].cells[j], header_fill)
        row_idx = 1
    for i,row in enumerate(data):
        for j,val in enumerate(row):
            set_cell_text(table.rows[i+row_idx].cells[j], '' if val is None else val, style='Small Table' if style_small else 'Normal', size=8.2 if style_small else 10)
    if col_widths:
        for row in table.rows:
            for cell, width in zip(row.cells, col_widths):
                cell.width = Mm(width)
    return table

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='Normal')
    p.style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p.paragraph_format.space_after = Pt(0)
    p.add_run(text)

def add_number(text, level=0):
    p = doc.add_paragraph(style='Normal')
    p.style = 'List Number' if level == 0 else 'List Number 2'
    p.paragraph_format.space_after = Pt(0)
    p.add_run(text)

def add_callout(title, text, fill='EEF5FB'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    cell = table.cell(0,0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.style = styles['Callout']
    r1 = p.add_run(title + ' ')
    r1.bold = True
    r2 = p.add_run(text)
    for r in (r1,r2):
        r.font.name='Calibri'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        r.font.size = Pt(10)
    doc.add_paragraph('')

# Header/footer with page number maybe simple text
for sec in doc.sections:
    header_p = sec.header.paragraphs[0]
    header_p.text = 'RADSAT-2U · Guía detallada de modelado térmico en ESATAN-TMS'
    header_p.style = styles['Small Table']
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Cover page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('RADSAT-2U')
r.bold = True; r.font.size = Pt(26); r.font.name='Calibri'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Calibri')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Guía detallada para crear un satélite térmico simplificado en ESATAN-TMS')
r.font.size = Pt(16); r.font.name='Calibri'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Calibri')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Documento de apoyo para TFG')
r.italic = True; r.font.size = Pt(12)

for _ in range(2):
    doc.add_paragraph('')
doc.add_picture('/mnt/data/radsat_schematic.png', width=Mm(165))
cap = doc.add_paragraph('Figura 1. Esquema simplificado del satélite propuesto, con vistas de distribución interna y secciones principales.')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.style = styles['Small Table']
doc.add_paragraph('')
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Versión regenerada · Marzo de 2026').bold = True
doc.add_paragraph('').alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Este documento convierte la propuesta del Excel en una memoria extensa y operativa.\n').italic = True
p.add_run('El objetivo es que puedas usarlo como guía de modelado, de redacción y de justificación de diseño en tu TFG.')
doc.add_page_break()

# Index / roadmap
doc.add_heading('Cómo usar este documento', level=1)
doc.add_paragraph(
    'Este texto está redactado para que puedas utilizarlo de tres formas distintas: '
    'como guía de diseño del satélite, como procedimiento de trabajo dentro de ESATAN-TMS '
    'y como base para redactar la memoria de tu TFG. Todo lo que aparece como “propuesta” '
    'es una decisión de diseño pensada para que el modelo sea asumible y no excesivamente complejo.'
)
add_bullet('Si necesitas avanzar rápido, empieza por las secciones 2, 3, 6 y 8.')
add_bullet('Si tu prioridad es reproducir el modelo en ESATAN-TMS, céntrate en la sección 9.')
add_bullet('Si estás escribiendo la memoria, las secciones 1, 4, 5, 7 y 10 te sirven para justificar decisiones.')
add_bullet('Los anexos del final condensan las tablas operativas: jerarquía, nodos, materiales, ópticas, potencias y checklist.')
add_callout('Idea central:', 'el modelo propuesto no intenta ser el satélite más realista posible, sino un satélite térmicamente coherente, suficientemente distinto del trabajo de referencia y lo bastante simple como para que puedas defenderlo con claridad.')

# 1 Alcance
doc.add_heading('1. Alcance, objetivo y filosofía del modelo', level=1)
doc.add_paragraph(
    'El satélite propuesto es un CubeSat 2U para monitorización de radiación en órbita baja, '
    'con una carga útil sencilla tipo dosímetro y una arquitectura térmica basada en estructura de aluminio, '
    'dos paneles solares laterales, una superficie radiadora dedicada, aislamiento multicapa y un conjunto de cajas electrónicas disipativas. '
    'La misión concreta puede formularse como “monitorización de radiación y housekeeping térmico-energético”, '
    'lo que permite justificar la existencia de la carga útil sin introducir elementos ópticos complejos.'
)
doc.add_paragraph(
    'La clave del planteamiento es mantener la lógica del ejemplo que te habían mostrado —jerarquía padre-hijo, '
    'componentes hasta nivel 3, materiales limitados y un número de nodos razonable—, pero cambiar la misión, '
    'la distribución interna y la naturaleza de la carga útil. De ese modo, el trabajo resultante es parecido en ambición '
    'y metodología, pero claramente distinto en contenido.'
)
doc.add_paragraph(
    'Para el modelo térmico se ha escogido una discretización de aproximadamente 102 nodos térmicos equivalentes, '
    'que es una cifra muy manejable para un TFG. Es suficiente para distinguir caras, equipos, radiador, manta MLI y paneles solares, '
    'pero no tan alta como para que el trabajo se convierta en una campaña de depuración de nodos y conductores.'
)
add_callout('Qué estás modelando realmente:', 'un modelo térmico lumped-parameter con cierta ayuda geométrica. No estás intentando reproducir cada tornillo, cada soporte ni cada lámina física del MLI; estás creando una representación térmica eficaz, justificable y trazable.')

# 2 Concepto de misión
doc.add_heading('2. Definición del satélite propuesto', level=1)
doc.add_heading('2.1. Concepto de misión', level=2)
doc.add_paragraph(
    'La opción elegida es un CubeSat 2U con un dosímetro o sensor de radiación sencillo como carga útil principal. '
    'En comparación con una misión óptica, esta alternativa tiene varias ventajas para un TFG térmico: elimina la óptica '
    'de precisión, evita tener que justificar una apertura o una lente compleja, reduce la sensibilidad del modelo a alineamientos '
    'geométricos finos y permite tratar la carga útil como un bloque sensor más una pequeña caja electrónica asociada.'
)
doc.add_heading('2.2. Razones por las que esta opción es adecuada', level=2)
for t in [
    'Es distinta del ejemplo de referencia porque no depende de una cámara ni de un barrido óptico.',
    'Es coherente con un satélite universitario o de baja complejidad.',
    'Facilita el trabajo térmico porque la carga útil puede modelarse como volúmenes disipativos compactos.',
    'Permite introducir una superficie radiadora clara y justificar de forma intuitiva la MLI y los paneles solares.',
    'Ayuda a que la memoria del TFG sea comprensible: la arquitectura se explica rápido y cada decisión tiene sentido.'
]:
    add_bullet(t)
doc.add_heading('2.3. Subsistemas principales', level=2)
doc.add_paragraph(
    'La arquitectura del modelo se divide en seis familias: estructura, sistema de potencia, aviónica, comunicaciones, '
    'carga útil y control térmico, a las que se añade el conjunto de paneles solares. Esta división es especialmente útil '
    'en ESATAN-TMS porque permite asignar rangos de nodos claros, revisar submodelos por separado y construir una jerarquía limpia.'
)

# 3 Architecture
doc.add_heading('3. Arquitectura jerárquica hasta nivel 3', level=1)
doc.add_paragraph(
    'La jerarquía es uno de los puntos que más conviene dejar cerrados desde el principio. '
    'En un TFG, una jerarquía consistente vale casi tanto como una geometría bonita, porque evita errores de etiquetado, '
    'de asignación de materiales y de interpretación posterior.'
)
# compact hierarchy table
hier_table = [
    ['RADSAT-2U', '0', 'None', '10000-77999', '102', 'Modelo completo'],
    ['STRUCTURE', '1', 'RADSAT-2U', '10000-15999', '28', 'Cuerpo + bandeja'],
    ['POWER_SYS', '1', 'RADSAT-2U', '20000-21999', '12', 'Batería + EPS'],
    ['AVIONICS', '1', 'RADSAT-2U', '30000-31999', '8', 'OBC + ADCS'],
    ['COMMS', '1', 'RADSAT-2U', '40000-41499', '6', 'Radio + soporte antena'],
    ['PAYLOAD', '1', 'RADSAT-2U', '50000-50999', '8', 'Dosímetro y electrónica'],
    ['THERMAL_CTRL', '1', 'RADSAT-2U', '60000-66999', '28', 'Radiador + MLI'],
    ['SOLAR_ARRAY', '1', 'RADSAT-2U', '70000-71999', '12', 'Dos paneles laterales'],
]
add_table(hier_table, headers=['Componente', 'Nivel', 'Padre', 'Rango', 'Nodos', 'Descripción'], style_small=True)
doc.add_paragraph('')
doc.add_paragraph(
    'A nivel 2, el detalle mínimo recomendado es: BODY e INTERNAL_TRAY dentro de estructura; BATTERY y EPS dentro de potencia; '
    'OBC y ADCS en aviónica; RADIO_UHF y ANTENNA en comunicaciones; DOSIMETER dentro de carga útil; '
    'RADIATOR y MLI dentro de control térmico; y PANEL_LEFT y PANEL_RIGHT dentro del conjunto solar. '
    'A nivel 3, lo más operativo es trabajar con caras o cajas térmicas equivalentes, como BODY_POS_X, BODY_NEG_X, '
    'TRAY_MAIN, BATTERY_PACK, EPS_BOARD, OBC_BOARD, ADCS_BOARD, TRX_MODULE, DOSIMETER_HEAD, DOSIMETER_ELEC, '
    'RADIATOR_NADIR, MLI_FACE_i, SP_LEFT y SP_RIGHT.'
)
add_callout('Regla práctica:', 'si dudas entre crear un componente nuevo o absorberlo dentro de otro, en un TFG suele ser mejor absorberlo, salvo que tenga una función térmica muy clara o una disipación relevante.')

# 4 geometry
doc.add_heading('4. Geometría, dimensiones y distribución física', level=1)
doc.add_heading('4.1. Sistema de referencia', level=2)
doc.add_paragraph(
    'Se adopta un sistema de referencia centrado en el centro geométrico del cuerpo del satélite. '
    'La dimensión larga del 2U se sitúa sobre el eje Z. Por tanto, el cuerpo ocupa aproximadamente '
    'x ∈ [-50, +50] mm, y ∈ [-50, +50] mm y z ∈ [-110, +110] mm.'
)
doc.add_heading('4.2. Dimensiones recomendadas', level=2)
dim_table = [
    ['Cuerpo exterior', 'Caja 2U', '100 × 100 × 220 mm', '2 mm', 'Al_6061'],
    ['Bandeja interna', 'Placa', '96 × 96 mm', '2 mm', 'Al_6061'],
    ['Radiador', 'Placa', '90 × 90 mm', '2 mm', 'Al_6061 + OSR'],
    ['Panel solar izquierdo', 'Placa', '95 × 210 mm', '2 mm', 'GaAs + vidrio'],
    ['Panel solar derecho', 'Placa', '95 × 210 mm', '2 mm', 'GaAs + vidrio'],
    ['Battery pack', 'Bloque', '90 × 45 × 15 mm', '—', 'Al_6061 equivalente'],
    ['EPS', 'Caja/placa', '90 × 90 × 8 mm', '—', 'Al_6061 equivalente'],
    ['OBC', 'Caja/placa', '90 × 90 × 8 mm', '—', 'Al_6061 equivalente'],
    ['ADCS', 'Caja/placa', '90 × 90 × 8 mm', '—', 'Al_6061 equivalente'],
    ['TRX', 'Módulo RF', '90 × 60 × 12 mm', '—', 'Al_6061 equivalente'],
    ['Dosimeter head', 'Bloque sensor', '50 × 50 × 20 mm', '—', 'Al_6061 equivalente'],
    ['Dosimeter elec', 'Caja', '60 × 40 × 10 mm', '—', 'Al_6061 equivalente'],
]
add_table(dim_table, headers=['Elemento', 'Tipo', 'Dimensiones principales', 'Espesor', 'Material base'], style_small=True)
doc.add_heading('4.3. Distribución interna sugerida', level=2)
doc.add_paragraph(
    'Una distribución simple y defendible consiste en apilar los equipos principales a lo largo del eje Z, '
    'sobre o alrededor de una bandeja central. La batería puede ir en la zona inferior; por encima, EPS, OBC y ADCS; '
    'el transceptor ligeramente lateral; y en la zona superior la carga útil, separando el cabezal sensor de su electrónica. '
    'El radiador se reserva para la cara -Z y los paneles solares se colocan en las caras ±Y.'
)
placement = [
    ['BATTERY_PACK', '(0, 0, -45)', 'Centrado y cercano a la mitad inferior'],
    ['EPS_BOARD', '(0, 0, -10)', 'Sobre la bandeja, próximo a batería'],
    ['OBC_BOARD', '(0, 0, +20)', 'Centro del volumen interno'],
    ['ADCS_BOARD', '(0, 0, +50)', 'Sobre OBC, mismo criterio de montaje'],
    ['TRX_MODULE', '(+25, 0, +78)', 'Desplazado hacia una pared lateral'],
    ['DOSIMETER_ELEC', '(0, 0, +85)', 'Electrónica asociada a la carga útil'],
    ['DOSIMETER_HEAD', '(0, 0, +98)', 'En el extremo superior del satélite'],
]
add_table(placement, headers=['Equipo', 'Centro aproximado [mm]', 'Comentario'], style_small=True)
doc.add_paragraph(
    'Estas coordenadas no deben interpretarse como un requisito absoluto; son una propuesta para que la geometría sea consistente, '
    'fácil de dibujar y térmicamente interpretable. En la memoria del TFG puedes presentarlas como “posiciones nominales de modelado”.'
)
doc.add_picture('/mnt/data/esatan_flow.png', width=Mm(165))
cap = doc.add_paragraph('Figura 2. Flujo de trabajo recomendado dentro de ESATAN-TMS para este TFG.')
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.style = styles['Small Table']

# 5 materials optics
doc.add_heading('5. Materiales y propiedades termo-ópticas', level=1)
doc.add_paragraph(
    'Para mantener el alcance controlado, se recomienda trabajar con un conjunto muy compacto de materiales y recubrimientos. '
    'Este criterio es importante: cuantas más familias de materiales introduzcas, más difícil será justificar cada parámetro, '
    'verificar consistencia y explicar resultados.'
)
mat_table = [[r[0], r[1], r[2], r[3], r[5]] for r in mat_d]
add_table(mat_table, headers=['Material', 'k [W/mK]', 'Cp [J/kgK]', 'ρ [kg/m³]', 'Uso'], style_small=True)
doc.add_paragraph('')
opt_table = [[r[0], r[1], r[5], r[10]] for r in opt_d if r[0] != 'IR lens']
add_table(opt_table, headers=['Recubrimiento', 'ε IR', 'α solar', 'Uso'], style_small=True)
doc.add_paragraph('')
doc.add_paragraph(
    'La asignación base recomendada es: aluminio anodizado para las caras estructurales exteriores, pintura negra para cajas internas '
    'y bandeja, OSR para la superficie radiadora, “solar cells” en la cara frontal de los paneles y Kapton ITO en la cara trasera. '
    'La MLI se representa con una propiedad efectiva MLI_int aplicada a mantas equivalentes por cara.'
)
add_callout('Importante:', 'en un modelo de TFG no es necesario describir cada capa física del panel solar ni del MLI. Es perfectamente aceptable usar una descripción equivalente, siempre que la expliques con honestidad.')
doc.add_heading('5.1. Asignación componente → material → óptica', level=2)
assign_table = [[r[0], r[1], r[2], r[3]] for r in asig_d]
add_table(assign_table, headers=['Componente o cara', 'Material', 'Óptica', 'Justificación'], style_small=True)

# 6 nodal strategy
doc.add_heading('6. Estrategia de nodado y rangos de numeración', level=1)
doc.add_paragraph(
    'El nodado es el corazón del trabajo. Una de las ventajas de plantear un satélite sencillo es que puedes asignar rangos amplios y muy legibles: '
    '10xxx para estructura, 20xxx para potencia, 30xxx para aviónica, 40xxx para comunicaciones, 50xxx para payload, 60xxx para control térmico '
    'y 70xxx para paneles solares. Este esquema facilita la depuración y además queda muy bien explicado en la memoria.'
)
doc.add_heading('6.1. Número total de nodos', level=2)
doc.add_paragraph(
    'La propuesta utiliza 102 nodos térmicos equivalentes. Ese número sale de sumar las 6 caras del cuerpo con 4 nodos por cara, la bandeja interna, '
    'los equipos disipativos, el radiador, la MLI y los dos paneles solares. Para un TFG, es una cifra equilibrada: suficientemente rica como para estudiar '
    'gradientes elementales, pero todavía sencilla de auditar.'
)
node_groups = [
    ['BODY_POS_X / NEG_X / POS_Y / NEG_Y / POS_Z / NEG_Z', '10000–10503', '24', '4 nodos por cara estructural'],
    ['TRAY_MAIN', '15000–15003', '4', 'Bandeja interna principal'],
    ['BATTERY_PACK', '20000–20007', '8', 'Volumen equivalente 2×2×2'],
    ['EPS_BOARD', '21000–21003', '4', 'Placa o caja equivalente'],
    ['OBC_BOARD', '30000–30003', '4', 'Placa o caja equivalente'],
    ['ADCS_BOARD', '31000–31003', '4', 'Placa o caja equivalente'],
    ['TRX_MODULE', '40000–40003', '4', 'Módulo RF'],
    ['ANTENNA_BASE', '41000–41001', '2', 'Soporte simplificado'],
    ['DOSIMETER_HEAD', '50000–50003', '4', 'Cabezal sensor'],
    ['DOSIMETER_ELEC', '50100–50103', '4', 'Electrónica asociada'],
    ['RADIATOR_NADIR', '60000–60005', '6', 'Superficie radiadora dedicada'],
    ['MLI_FACE_1 … MLI_FACE_6', '61000–66017', '18', 'Tres nodos por cara equivalente'],
    ['SP_LEFT / SP_RIGHT', '70000–71005', '12', 'Seis nodos por panel'],
]
add_table(node_groups, headers=['Grupo', 'Rango', 'Nodos', 'Observación'], style_small=True)
doc.add_heading('6.2. Criterios para decidir cuántos nodos usar', level=2)
for t in [
    'Cuatro nodos bastan para una placa o caja pequeña si no buscas gradientes finos.',
    'Seis nodos en cada panel solar permiten reproducir el criterio del ejemplo sin complicar demasiado el modelo.',
    'Ocho nodos para batería son útiles porque ese elemento suele ser térmicamente sensible y conviene no dejarlo con un único nodo.',
    'Tres nodos por manta MLI equivalente son más que suficientes en un modelo conceptual.',
    'Si una pieza no tiene disipación propia y solo transmite calor, probablemente no necesita muchos nodos.'
]:
    add_bullet(t)

# 7 contacts
doc.add_heading('7. Contactos, acoplos y conductores propuestos', level=1)
doc.add_paragraph(
    'Ningún modelo térmico funciona bien si la geometría está definida pero los contactos están mal planteados. '
    'En este diseño, la filosofía es clara: los equipos internos acoplan principalmente a la bandeja interna y, a través de ella, a la estructura. '
    'Además, existe acoplo radiativo interno simplificado entre cajas, bandeja y superficies internas.'
)
doc.add_paragraph(
    'Como punto de partida para un TFG, puedes trabajar con conductancias lumped propuestas. '
    'No son “verdades absolutas”; son valores iniciales defendibles que después puedes ajustar en una pequeña sensibilidad.'
)
cond_table = [
    ['BATTERY_PACK ↔ TRAY_MAIN', '0.30', 'Anclaje moderado, contacto conservador'],
    ['EPS_BOARD ↔ TRAY_MAIN', '0.60', 'Electrónica de potencia bien atornillada'],
    ['OBC_BOARD ↔ TRAY_MAIN', '0.50', 'Caja/placa con buen apoyo'],
    ['ADCS_BOARD ↔ TRAY_MAIN', '0.40', 'Algo menos rígido que EPS'],
    ['TRX_MODULE ↔ BODY_POS_X', '0.25', 'Cercano a una pared lateral'],
    ['DOSIMETER_ELEC ↔ TRAY_MAIN', '0.25', 'Electrónica asociada al sensor'],
    ['DOSIMETER_HEAD ↔ DOSIMETER_ELEC', '0.10', 'Unión interna moderada'],
    ['TRAY_MAIN ↔ BODY faces', '1.50 total', 'Conducción estructural global'],
    ['PANEL_LEFT / RIGHT ↔ BODY_±Y', '0.20 cada uno', 'Unión mecánica simplificada'],
    ['RADIATOR_NADIR ↔ BODY_NEG_Z', '1.00', 'Contacto alto deseable'],
]
add_table(cond_table, headers=['Enlace', 'G [W/K] propuesta', 'Comentario'], style_small=True)
doc.add_paragraph(
    'Si prefieres trabajar con conductancia de contacto por área, puedes transformar estos valores en un producto h·A equivalente, '
    'pero para un TFG suele ser más transparente hablar directamente de conductores lumped entre nodos o grupos de nodos.'
)

# 8 powers cases
doc.add_heading('8. Potencias internas y casos térmicos', level=1)
doc.add_heading('8.1. Potencias internas propuestas', level=2)
pot_table = [[r[0], r[1], r[2], r[3], r[5]] for r in pot_d]
add_table(pot_table, headers=['Elemento', 'Nominal [W]', 'Cold [W]', 'Hot [W]', 'Comentario'], style_small=True)
doc.add_heading('8.2. Casos térmicos mínimos a ejecutar', level=2)
cases = [
    ['NOMINAL', 'Disipación nominal con entorno representativo', 'Comportamiento base del satélite'],
    ['HOT', 'Máxima disipación, Sol/albedo/IR desfavorables', 'Comprobar T_max de radio, OBC, payload'],
    ['COLD', 'Baja disipación y eclipse o entorno frío equivalente', 'Comprobar T_min de batería y carga útil'],
    ['SENSIBILIDAD', 'Variación de 10–20 % en conductancias o propiedades', 'Mostrar robustez del modelo'],
]
add_table(cases, headers=['Caso', 'Descripción', 'Objetivo'], style_small=True)
doc.add_paragraph(
    'En la memoria conviene dejar claro que, aunque el satélite sea sencillo, sí se han analizado al menos un caso caliente y un caso frío. '
    'Eso refuerza mucho la validez del trabajo. Si tu tutor quiere más profundidad, puedes añadir un barrido de órbita, de beta angle '
    'o una sensibilidad al recubrimiento del radiador.'
)
doc.add_heading('8.3. Órbita de referencia recomendada', level=2)
doc.add_paragraph(
    'Como órbita base para el TFG puedes adoptar una LEO circular de 500–600 km. '
    'Una elección razonable es 550 km con actitud fija al cuerpo. No necesitas complicarte con cinemáticas avanzadas a menos que tu tutor lo pida. '
    'Lo importante es definir con claridad qué cara recibe el Sol, cuál ve mejor el espacio profundo y cómo cambia la carga térmica en las fases de sol y eclipse.'
)
add_callout('Consejo para defender el modelo:', 'es mejor presentar una órbita simple y muy bien explicada que una misión compleja poco controlada. La simplicidad, cuando está justificada, juega a tu favor.')

# 9 step by step ESATAN
doc.add_heading('9. Procedimiento detallado en ESATAN-TMS', level=1)
doc.add_paragraph(
    'La secuencia de trabajo recomendada sigue la lógica de los módulos oficiales de ESATAN-TMS. '
    'La información pública del fabricante describe Workbench como el entorno integrado de pre y postproceso, '
    'Mission como el módulo de entorno orbital, Radiative como el entorno de cálculo radiativo, Thermal como el solver térmico '
    'y ThermNV como la herramienta de inspección y postproceso de la red térmica. Esa división encaja muy bien con el tipo de modelo que necesitas [R1–R5].'
)
steps_intro = [
    'Crea un proyecto limpio y define desde el inicio las unidades que vas a usar.',
    'Construye primero la geometría y la jerarquía, no al revés.',
    'Da de alta materiales y ópticas antes de empezar a enlazar casos.',
    'Congela una estrategia de nodos y rangos antes de lanzar cálculos.',
    'Valida la red y la trazabilidad del modelo antes de entrar a afinar resultados.'
]
for t in steps_intro:
    add_number(t)

doc.add_heading('9.1. Crear el proyecto base', level=2)
doc.add_paragraph(
    'Abre Workbench y crea un proyecto nuevo. Guarda el proyecto en una carpeta exclusiva para el TFG; '
    'evita mezclarlo con otros ensayos. Define desde el principio el sistema de unidades principales: milímetros para geometría, '
    'vatios para potencias y kelvin o grados Celsius para la presentación térmica, según prefieras trabajar. '
    'El objetivo de este primer paso es dejar preparado un contenedor limpio y repetible.'
)
doc.add_heading('9.2. Definir variables globales', level=2)
doc.add_paragraph(
    'Antes de dibujar la geometría, crea variables globales. Por ejemplo: L_body = 220 mm, W_body = 100 mm, H_body = 100 mm, '
    't_wall = 2 mm, L_panel = 210 mm, W_panel = 95 mm, L_rad = 90 mm y W_rad = 90 mm. '
    'Esto no solo hace más limpio el modelo; también te permite justificar que la geometría es paramétrica y que podrías hacer sensibilidades sencillas.'
)
doc.add_heading('9.3. Crear la estructura', level=2)
doc.add_paragraph(
    'Modela el cuerpo principal como una caja hueca o, si prefieres simplificar aún más, como seis placas equivalentes. '
    'Para un TFG, la segunda opción suele bastar y, además, hace muy legible la numeración por caras. '
    'Añade la bandeja interna como una placa aproximadamente a media altura del satélite. '
    'No hace falta modelar railes, tornillería ni rigidizadores, salvo que quieras usarlos como conductores equivalentes.'
)
doc.add_heading('9.4. Añadir los equipos internos', level=2)
doc.add_paragraph(
    'Crea volúmenes o placas equivalentes para BATTERY_PACK, EPS_BOARD, OBC_BOARD, ADCS_BOARD, TRX_MODULE, DOSIMETER_HEAD y DOSIMETER_ELEC. '
    'Sitúalos según la tabla de distribución propuesta. Si el entorno de tu versión lo hace más cómodo, puedes modelarlos como cajas o como superficies '
    'con una capacidad térmica equivalente; lo importante es que cada equipo tenga masa térmica, un conjunto de nodos razonable y un camino claro de evacuación.'
)
doc.add_heading('9.5. Añadir paneles, radiador y MLI', level=2)
doc.add_paragraph(
    'Añade dos paneles solares fijos en las caras ±Y. Mantenerlos fijos simplifica enormemente el problema. '
    'Define el radiador como una placa sobre la cara -Z. La MLI no la modeles manta por manta; crea superficies o propiedades equivalentes por cara. '
    'En versiones recientes, la documentación pública de ESATAN-TMS indica incluso una ampliación del soporte para definiciones multicapa de aislamiento, '
    'lo que refuerza esta forma de trabajo [R6].'
)
doc.add_heading('9.6. Definir materiales y propiedades ópticas', level=2)
doc.add_paragraph(
    'Carga los cuatro materiales base y las propiedades ópticas elegidas. Después, asigna a cada pieza su material y su recubrimiento. '
    'Haz esta comprobación conscientemente: una gran parte de los errores de un modelo térmico sencillo no viene de la geometría, sino de un recubrimiento mal asignado '
    'o de una pieza a la que se olvidó dar capacidad térmica.'
)
doc.add_heading('9.7. Establecer nodos y rangos', level=2)
doc.add_paragraph(
    'Define el esquema de numeración antes de dejar que el programa genere o reordene elementos sin control. '
    'Usa rangos coherentes, como los de este documento, y vincula cada grupo a su etiqueta correspondiente. '
    'Si la herramienta te permite fijar nodos por superficie o por sólido, aprovecha esa posibilidad para mantener el control del modelo y no acabar con un mallado innecesario.'
)
doc.add_heading('9.8. Definir conductores y acoplos', level=2)
doc.add_paragraph(
    'Conecta cada equipo interno a la bandeja y/o a la estructura. Después, conecta la bandeja con las caras del cuerpo y el radiador con su soporte estructural. '
    'Por último, establece el acoplo de los paneles con las caras laterales. '
    'Si dudas en una primera iteración, es preferible sobredocumentar qué conductor has añadido y por qué, antes que tener una red opaca.'
)
doc.add_heading('9.9. Montar el caso radiativo y orbital', level=2)
doc.add_paragraph(
    'Una vez lista la geometría, prepara el problema radiativo. La información pública oficial describe Radiative como el entorno de cálculo '
    'de factores de forma e intercambio radiativo y Mission como el módulo orientado a entorno orbital y cargas externas [R3–R4]. '
    'Para un TFG básico, basta con definir la orientación del satélite, la órbita de referencia, el albedo y el IR planetario de manera coherente, '
    'sin forzar una complejidad mayor de la necesaria.'
)
doc.add_heading('9.10. Definir casos térmicos y solver', level=2)
doc.add_paragraph(
    'Crea al menos los casos NOMINAL, HOT y COLD. Asigna en cada uno la columna de potencias correspondiente y verifica que la misma geometría se conserva. '
    'Lanza primero un estacionario para detectar errores groseros y, cuando el modelo sea estable, añade un transitorio simple si tu tutor lo considera útil.'
)
doc.add_heading('9.11. Validar la red en ThermNV', level=2)
doc.add_paragraph(
    'La documentación oficial describe ThermNV como herramienta de visualización, inspección y reporting de la red térmica [R5]. '
    'Úsalo para revisar nodos huérfanos, enlaces incoherentes, subredes aisladas y resultados extremos poco verosímiles. '
    'Es una parte muy valiosa del trabajo, porque demuestra que no te limitaste a lanzar el solver, sino que también auditaste el modelo.'
)
doc.add_heading('9.12. Guardar, documentar y congelar versiones', level=2)
doc.add_paragraph(
    'Cuando una iteración del modelo funcione, congélala con un nombre claro: por ejemplo, v01_base, v02_optics, v03_hotcase. '
    'Haz capturas, exporta tablas de nodos y deja rastro de cada decisión importante. Esa disciplina te ahorrará tiempo cuando redactes la memoria final.'
)

doc.add_heading('10. Qué deberías enseñar en la memoria del TFG', level=1)
for t in [
    'Una figura general del satélite con orientación y componentes principales.',
    'La jerarquía hasta nivel 3 con rangos de nodos.',
    'La tabla de materiales y la tabla de recubrimientos.',
    'Las potencias internas y los casos analizados.',
    'Una explicación honesta de los supuestos de simplificación.',
    'Temperaturas máximas y mínimas de batería, carga útil, radio y estructura.',
    'Una breve sensibilidad o, al menos, una discusión sobre incertidumbres.'
]:
    add_bullet(t)
doc.add_paragraph(
    'En otras palabras, tu memoria no necesita demostrar que el satélite es industrialmente perfecto; '
    'necesita demostrar que tu modelo es coherente, controlado y suficientemente profundo para responder preguntas térmicas relevantes.'
)

doc.add_heading('11. Errores comunes y cómo evitarlos', level=1)
errors = [
    ('Modelar demasiados detalles desde el principio', 'Empieza por una versión limpia y mínima; después, añade complejidad solo si aporta valor térmico.'),
    ('Usar demasiados materiales', 'Cuantos menos materiales y recubrimientos, más fácil será auditar el modelo.'),
    ('No controlar los rangos de nodos', 'Define numeración por subsistemas desde el primer día.'),
    ('Olvidar justificar simplificaciones', 'Explica siempre por qué una caja o una manta se ha representado de forma equivalente.'),
    ('Centrarse solo en temperaturas finales', 'Revisa también conductores, balance radiativo y rutas de calor.'),
    ('No distinguir dato de referencia y propuesta propia', 'En la memoria separa claramente qué heredaste como estilo metodológico y qué diseñaste tú.')
]
add_table(errors, headers=['Error', 'Cómo evitarlo'], style_small=True)

doc.add_heading('12. Plan de trabajo recomendado', level=1)
schedule = [
    ['Fase 1', 'Definir misión, jerarquía, dimensiones y tabla de materiales'],
    ['Fase 2', 'Construir geometría y asignar materiales/ópticas'],
    ['Fase 3', 'Fijar nodos, rangos y conductores'],
    ['Fase 4', 'Montar casos térmicos HOT/COLD/NOMINAL'],
    ['Fase 5', 'Validar resultados y preparar figuras/tablas para memoria'],
    ['Fase 6', 'Redactar justificación, resultados y discusión final'],
]
add_table(schedule, headers=['Fase', 'Objetivo'], style_small=True)
doc.add_paragraph(
    'Si trabajas de esta manera, siempre tendrás un modelo funcional en cada etapa. Eso es muy importante en un TFG: '
    'evita quedarte atrapado en un modelo “casi terminado” que todavía no resuelve nada.'
)

# Appendices
doc.add_page_break()
doc.add_heading('Anexo A. Tabla operativa de jerarquía', level=1)
jer_compact = [[r[0], r[1], r[2], r[3], r[6], r[10]] for r in jer_d]
add_table(jer_compact, headers=['Componente', 'Nivel', 'Padre', 'Rango', 'Nodos', 'Nota'], style_small=True)

doc.add_heading('Anexo B. Tabla operativa de nodos nivel 3', level=1)
nod_compact = [[r[0], r[1], r[2], r[3], r[4], r[5]] for r in nod_d]
add_table(nod_compact, headers=['Elemento', 'Nivel', 'Padre', 'Rango', 'Nodos', 'Propósito'], style_small=True)

doc.add_heading('Anexo C. Dimensiones propuestas', level=1)
dim_compact = [[r[0], r[1], f"{r[2]}×{r[3]}×{r[4]} {r[6]}", r[7], r[8]] for r in dim_d]
add_table(dim_compact, headers=['Componente', 'Geometría', 'Dimensiones', 'Material', 'Nota'], style_small=True)

doc.add_heading('Anexo D. Materiales y ópticas', level=1)
add_table([[r[0], r[1], r[2], r[3], r[5]] for r in mat_d],
          headers=['Material', 'k', 'Cp', 'Densidad', 'Uso'], style_small=True)
doc.add_paragraph('')
add_table([[r[0], r[1], r[5], r[10]] for r in opt_d],
          headers=['Óptica', 'ε IR', 'α solar', 'Uso'], style_small=True)

doc.add_heading('Anexo E. Potencias y checklist', level=1)
add_table([[r[0], r[1], r[2], r[3], r[5]] for r in pot_d],
          headers=['Elemento', 'Nominal', 'Cold', 'Hot', 'Comentario'], style_small=True)
doc.add_paragraph('')
add_table([[r[0], r[1]] for r in check_d],
          headers=['Tema', 'Comprobación'], style_small=True)

doc.add_heading('Anexo F. Referencias y notas finales', level=1)
doc.add_paragraph(
    'Referencias oficiales ESATAN-TMS consultadas para la redacción del flujo de trabajo y la descripción de módulos:'
)
refs = [
    '[R1] ESATAN-TMS Workbench. Página oficial de producto. Descripción de Workbench como entorno integrado de pre y postproceso, modelado geométrico y control de casos.',
    '[R2] ESATAN-TMS Mission. Página oficial de producto. Descripción de Mission como entorno para análisis térmico orbital.',
    '[R3] ESATAN-TMS Radiative. Página oficial de producto. Descripción del cálculo radiativo y de factores de intercambio.',
    '[R4] ESATAN-TMS Home / Products overview. Descripción general del flujo completo desde geometría hasta postproceso.',
    '[R5] ESATAN-TMS ThermNV. Página oficial de producto. Descripción de validación, inspección y reporting de la red térmica.',
    '[R6] ESATAN-TMS 2026 release news. Información pública sobre nuevas funciones en Workbench y definición multicapa de aislamiento.'
]
for ref in refs:
    doc.add_paragraph(ref, style='Small Table')
doc.add_paragraph('')
doc.add_paragraph(
    'Nota final: este documento mezcla criterios heredados del estilo del ejemplo que se te facilitó '
    '(por ejemplo, jerarquía, materiales compactos y control de rangos) con una propuesta propia de misión, '
    'geometría, nodado, ubicación de equipos y casos térmicos. Es decir, es una base original y usable para tu TFG.'
)

# save
docx_path = '/mnt/data/RADSAT_2U_Guia_Detallada_ESATAN_TFG_v2.docx'
doc.save(docx_path)
docx_path, os.path.getsize(docx_path)

import subprocess, os, textwrap, json, glob, sys, pathlib, shlex, tempfile
docx_path = '/mnt/data/RADSAT_2U_Guia_Detallada_ESATAN_TFG_v2.docx'
out_dir = '/mnt/data/radsat_v2_render'
os.makedirs(out_dir, exist_ok=True)
cmd = ['python','/home/oai/skills/docx/render_docx.py', docx_path, '--output_dir', out_dir, '--emit_pdf']
print('running', cmd)
res = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
print('returncode', res.returncode)
print('stdout', res.stdout[:3000])
print('stderr', res.stderr[:3000])
print('files', glob.glob(out_dir+'/*')[:20], 'count', len(glob.glob(out_dir+'/*')))

from PIL import Image, ImageOps, ImageDraw
import glob, os, math
imgs = [Image.open(p).convert('RGB') for p in sorted(glob.glob('/mnt/data/radsat_v2_render/page-*.png'))]
# create contact sheet
thumb_w = 240
thumbs = []
for i, im in enumerate(imgs, start=1):
    ratio = thumb_w / im.width
    th = int(im.height * ratio)
    t = im.resize((thumb_w, th))
    canvas = Image.new('RGB', (thumb_w, th+22), 'white')
    canvas.paste(t, (0,22))
    d = ImageDraw.Draw(canvas)
    d.text((6,4), f'Pág. {i}', fill='black')
    thumbs.append(canvas)
cols = 3
rows = math.ceil(len(thumbs)/cols)
cell_w = max(t.width for t in thumbs)+10
cell_h = max(t.height for t in thumbs)+10
sheet = Image.new('RGB', (cols*cell_w, rows*cell_h), '#DDDDDD')
for idx, t in enumerate(thumbs):
    x = (idx%cols)*cell_w + 5
    y = (idx//cols)*cell_h + 5
    sheet.paste(t,(x,y))
contact='/mnt/data/radsat_v2_contact.png'
sheet.save(contact)
contact

import re, glob, os
pages = sorted(glob.glob('/mnt/data/radsat_v2_render/page-*.png'), key=lambda p:int(re.search(r'page-(\d+)\.png',p).group(1)))
print([os.path.basename(p) for p in pages])

from PIL import Image, ImageOps, ImageDraw
import glob, os, math, re
imgs = [Image.open(p).convert('RGB') for p in sorted(glob.glob('/mnt/data/radsat_v2_render/page-*.png'), key=lambda p:int(re.search(r'page-(\d+)\.png',p).group(1)))]
thumb_w = 240
thumbs=[]
for i,im in enumerate(imgs, start=1):
    ratio = thumb_w / im.width
    th = int(im.height * ratio)
    t=im.resize((thumb_w,th))
    canvas=Image.new('RGB',(thumb_w, th+22),'white')
    canvas.paste(t,(0,22))
    d=ImageDraw.Draw(canvas)
    d.text((6,4),f'Pág. {i}',fill='black')
    thumbs.append(canvas)
cols=3
rows=math.ceil(len(thumbs)/cols)
cell_w=max(t.width)

import shutil, os, glob
src_docx='/mnt/data/RADSAT_2U_Guia_Detallada_ESATAN_TFG_v2.docx'
src_pdf='/mnt/data/radsat_v2_render/RADSAT_2U_Guia_Detallada_ESATAN_TFG_v2.pdf'
final_docx='/mnt/data/RADSAT_2U_Guia_Detallada_ESATAN_TFG_regenerado.docx'
final_pdf='/mnt/data/RADSAT_2U_Guia_Detallada_ESATAN_TFG_regenerado.pdf'
shutil.copyfile(src_docx, final_docx)
shutil.copyfile(src_pdf, final_pdf)
print(final_docx, os.path.getsize(final_docx))
print(final_pdf, os.path.getsize(final_pdf))
